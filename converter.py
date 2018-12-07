#!/usr/local/bin/python
# -*- coding: utf-8 -*-

########## LIBRARIES ##########
from docx import Document
from docx.table import Table
from lxml import etree
from datetime import date
import sys, os, zipfile, re
import styleElements as styEl
import argparse

########## GLOBAL VARIABLES ##########
tableOpen = listOpen = lgOpen = citOpen = nestedCitOpen = speechOpen = nestedSpeechOpen = False
bodyOpen = backOpen = frontOpen = False
global_header_level = 0
isDir = False
inDocument = False
global_list_level = 0
footnoteNum = endnoteNum = 0
idTracker = []
titlePage = None
footnotes = []
endnotes = []
root = None
debugme = False
unsupported_char = {}
badheader_text = []
metaTemplate = 'teiHeader.dat'
convert_options = None


########## FUNCTIONS ##########

def doFootEndnotes(inputFile):
    """
    Parses Word XML footnotes and endnotes into global lists to be referenced during conversion
    """
    global footnotes, endnotes

    # note:	the global footnote and endnote lists are 0 indexed (e.g. footnotes[0] is the footnote #1)

    document = zipfile.ZipFile(inputFile)

    # write content of footnotes.xml into footnotes[]

    xml_content = document.read('word/footnotes.xml')
    root = etree.fromstring(xml_content)

    curIndex = 0

    foot = root.findall('w:footnote', root.nsmap)
    for f in foot:
        if curIndex > 1:
            text = f.findall('.//w:t', root.nsmap)
            s = ""
            for t in text:
                s += t.text
            footnotes.append(s)
        curIndex += 1

    # write content of endnotes.xml into endnotes[]

    xml_content = document.read('word/endnotes.xml')
    root = etree.fromstring(xml_content)

    curIndex = 0

    end = root.findall('w:endnote', root.nsmap)
    for f in end:
        if curIndex > 1:
            text = f.findall('.//w:t', root.nsmap)
            s = ""
            for t in text:
                s += t.text
            endnotes.append(s)
        curIndex += 1
    document.close()


def doMetadata(metaTable):
    global metaTemplate
    # load metadata schema
    try:
        #if
        f = open(metaTemplate, 'rb')
        metaText = f.read()
    except:
        print "\t Error: teiHeader.dat file not in current working directory"
        print "\t Current directory: {0}".format(os.getcwd())
        sys.exit(1)


    # Fill out metadata matching on string in metatable with {strings} in template (teiHeader.dat)
    metaText = metaText.replace("{Digital Creation Date}", str(date.today()))
    problems_on = False
    for rwnum in range(0, len(metaTable.rows)):
        try:
            label = metaTable.cell(rwnum, 1).text.strip()
            if not label:
                label = metaTable.cell(rwnum, 0).text.strip()
            rowval = metaTable.cell(rwnum, 3).text.strip()
            # All Uppercase are Headers in the table skip
            if label.isupper():
                if label == 'PROBLEMS':
                    problems_on = True
                continue  # Skip labels
            if problems_on:
                if '<encodingDesc>' not in metaText:
                    metaText = metaText.replace(u'</fileDesc>', u'</fileDesc><encodingDesc><editorialDecl n="problems"><interpretation n="{}">'.format(label) +
                                            u'<p>{}</p>'.format(rowval) +
                                            u'</interpretation></editorialDecl></encodingDesc>')
                else:
                    metaText = metaText.replace(u'</interpretation></editorialDecl>', u'</interpretation><interpretation n="{}">'.format(label) +
                                            u'<p>{}</p>'.format(rowval) +
                                            u'</interpretation></editorialDecl>')
            temppt = label.split(' (')  # some rows have " (if applicable)" or possible some other instruction
            label = temppt[0]
            srclbl = "{" + label + "}"
            metaText = metaText.replace(srclbl, rowval)

        except IndexError as e:
            print "Index error in iterating metatable: {}".format(e)
        except TypeError as e:
            print "Type error in iterating metatable: {}".format(e)


    metaText = re.sub(r'\{([^\}]+)\}', r'<!--\1-->', metaText)
    metaText += "</text></TEI.2>"
    return metaText  # ignoring old code below

    # OLD METHOD fill metadata schema with data from metadata table in document (Deprecated function never gets this far)
    # BASIC METADATA
    metaText = metaText.replace("{Title of Text}", metaTable.cell(1, 3).text)
    metaText = metaText.replace("{Cover Page}", metaTable.cell(2, 3).text)

    temp = metaTable.cell(3, 3).text.encode('utf-8').splitlines()
    if len(temp) > 0:
        metaText = metaText.replace("{Cover Title Tib}", temp[0].decode('utf-8'))
    if len(temp) > 1:
        metaText = metaText.replace("{Cover Title San in Tib}", temp[1].decode('utf-8'))
    if len(temp) > 2:
        metaText = metaText.replace("{Cover Title San in Lanydza}", temp[2].decode('utf-8'))

    # String Labels in the XML and corresponding metadata table cell numbers (X, Y)
    # For inserting metadata into template. Could have different versions of this for different tables.


    # Old direct correspance
    metaRows = [
        ("{Title on Spine}", 4, 3),
        ("{Margin Title}", 5, 3),
        ("{Author of Text}", 6, 3),
        ("{Name of Collection}", 7, 3),
        ("{Publisher Name}", 9, 3),
        ("{Publisher Place}", 10, 3),
        ("{Publisher Date}", 11, 3),
        ("{ISBN}", 13, 3),
        ("{Library Call-number}", 14, 3),
        ("{Other ID number}", 15, 3),
        ("{Volume Letter}", 16, 3),
        ("{Volume Number}", 17, 3),
        ("{Pagination of Text}", 18, 3),
        ("{Pages Represented in this file}", 19, 3),
        ("{Name of Agent Creating Etext}", 21, 3),
        ("{Date Process Begun}", 22, 3),
        ("{Date Process Finished}", 23, 3),
        ("{Place of Process}", 24, 3),
        ("{Method of Process (OCR, input)}", 25, 3),
        ("{Name of Proofreader}", 27, 3),
        ("{Date Proof Began}", 28, 3),
        ("{Date Proof Finished}", 29, 3),
        ("{Place of Proof}", 30, 3),
        ("{Name of Markup-er}", 32, 3),
        ("{Date Markup Began}", 33, 3),
        ("{Date Markup Finished}", 34, 3),
        ("{Place of Markup}", 35, 3),
        ("{Name of Converter}", 37, 3),
        ("{Date Conversion Began}", 38, 3),
        ("{Date Conversion Finished}", 39, 3),
        ("{Place of Conversion}", 40, 3),
        ("{Problem Cell 1}", 42, 1),
        ("{Problem Cell 2}", 42, 3)
    ]


    for metaField in metaRows:
        try:
            metaText = metaText.replace(metaField[0], metaTable.cell(metaField[1], metaField[2]).text)
        except IndexError:
            label = metaField[0].replace('{', '“').replace('}', '”')
            print "Cannot locate metadata table cell for {0} ({1}, {2})".format(label, metaField[1], metaField[2])

    # DATE
    metaText = metaText.replace("{Digital Creation Date}", str(date.today()))
    metaText = re.sub(r'\{([^\}]+)\}', r'<!--\1-->', metaText)

    metaText += "</text></TEI.2>"

    return metaText


def closingComments(lastElement):
    global global_header_level
    while global_header_level > 0:
        lastElement = lastElement.getparent()
        lastElement.append(etree.Comment("Close of Level: " + str(global_header_level)))
        global_header_level -= 1


def doTitle(par, lastElement):
    global titlePage
    titlePage = etree.Element("titlePage")
    docTitle = etree.SubElement(titlePage, "docTitle")
    titlePart = etree.SubElement(docTitle, "titlePart")
    titlePart.text = par.text


def doHeaders(par, lastElement, root):
    global frontOpen, bodyOpen, backOpen, global_header_level, idTracker, badheader_text
    styName = par.style.name
    # print styName
    closeStyle(styName, lastElement)

    skippedHeader = False

    # front
    if "Heading 0 Front" == styName:
        # close out previous paragraphs
        # while global_header_level > 0:
        #     lastElement = lastElement.getparent()
        #     lastElement.append(etree.Comment("Close of Level: " + str(global_header_level)))
        #     global_header_level -= 1
        # lastElement = lastElement.getparent()
        lastElement = root.find('text')
        # create Front section
        front = etree.SubElement(lastElement, "front")
        front.set("id", "a")
        idTracker = [0]
        head = etree.SubElement(front, "head")
        iterateRange(par, head)
        # head.text = par.text
        frontOpen = True
        # add title at the top of Front
        if titlePage is not None:
            front.append(titlePage)
        return front

    # body
    elif "Heading 0 Body" == styName:
        # close out previous paragraphs
        while global_header_level > 0:
            lastElement = lastElement.getparent()
            lastElement.append(etree.Comment("Close of Level: " + str(global_header_level)))
            global_header_level -= 1
        lastElement = lastElement.getparent()
        lastElement = root.find('text')
        # create Body section
        body = etree.SubElement(lastElement, "body")
        body.set("id", "b")
        idTracker = [0]
        head = etree.SubElement(body, "head")
        iterateRange(par, head)
        # head.text = par.text
        frontOpen = False
        bodyOpen = True
        return body

    # back
    elif "Heading 0 Back" == styName:
        # close out previous paragraphs
        while global_header_level > 0:
            lastElement = lastElement.getparent()
            lastElement.append(etree.Comment("Close of Level: " + str(global_header_level)))
            global_header_level -= 1
        lastElement = lastElement.getparent()
        lastElement = root.find('text')
        # create Back section
        back = etree.SubElement(lastElement, "back")
        back.set("id", "c")
        idTracker = [0]
        head = etree.SubElement(back, "head")
        iterateRange(par, head)
        # head.text = par.text
        bodyOpen = False
        backOpen = True
        return back

    # do header divs within front/body/back
    elif "Heading" in styName or "Interstitial" in styName:
        # if no front/body/back, print warning
        if not frontOpen and not bodyOpen and not backOpen:
            badheader_text.append(par.text)

        if "Interstitial" in styName:
            headingNum = str(global_header_level)
            print "INterstitial heading num: {}".format(headingNum)
        else:
            headingNum = styName.split(" ")[1]

        if headingNum.isdigit():
            headingNum = int(headingNum)

            if headingNum - 1 > global_header_level:
                print "\t Warning (IMPROPER HEADER NESTING): Jumped from Heading " + str(
                    global_header_level) + " to Heading " + str(headingNum)
                print "\t\t Header text: " + par.text
                skippedHeader = True

            # push new nested level
            if headingNum > global_header_level:
                while headingNum != global_header_level:
                    # Calculate <div> id
                    curID = ""
                    if frontOpen:
                        curID = "a"
                    elif bodyOpen:
                        curID = "b"
                    elif backOpen:
                        curID = "c"

                    if global_header_level > len(idTracker) - 1:
                        idTracker.append(1)
                    else:
                        idTracker[global_header_level] += 1

                    curID += str(idTracker[0])
                    for i in range(1, global_header_level + 1):
                        curID += "." + str(idTracker[i])

                    # push new level
                    global_header_level += 1

                    # Create <div> on current (newly pushed) level
                    if lastElement.tag == 'lg':
                        lastElement = lastElement.getparent()
                    lastElement = etree.SubElement(lastElement, "div")
                    lastElement.set("n", str(global_header_level))
                    lastElement.set("id", curID)
                    head = etree.SubElement(lastElement, "head")
                    if not skippedHeader:
                        iterateRange(par, head)
                    # head.text = par.text
                    skippedHeader = False

                return lastElement

            # pop out a level or remain on same level
            else:
                # pop out as many levels as necessry
                while headingNum != global_header_level:
                    if lastElement is None:
                        print "\tNo last element when trying to close parent header levels {0}".format(
                            global_header_level)
                        lastElement = getNewLastElement()
                    else:
                        leparent = lastElement.getparent()
                        if leparent is not None:
                            lastElement = lastElement.getparent()

                    lastElement.append(etree.Comment("Close of Level: " + str(global_header_level)))
                    idTracker.pop()
                    global_header_level -= 1

                # Calculate <div> id
                curID = ""
                if frontOpen:
                    curID = "a"
                elif bodyOpen:
                    curID = "b"
                elif backOpen:
                    curID = "c"

                idTracker[global_header_level - 1] += 1

                curID += str(idTracker[0])
                for i in range(1, global_header_level):
                    curID += "." + str(idTracker[i])

                if lastElement is None:
                    print "\tNo last element when trying to close level {0}".format(global_header_level)
                    lastElement = getNewLastElement()
                else:
                    leparent = lastElement.getparent()
                    if leparent is not None:
                        lastElement = leparent

                lastElement.append(etree.Comment("Close of Level: " + str(global_header_level)))

                # create div on current (or newly popped) level
                lastElement = etree.SubElement(lastElement, "div")
                lastElement.set("n", str(global_header_level))
                lastElement.set("id", curID)
                head = etree.SubElement(lastElement, "head")
                iterateRange(par, head)
                # head.text = par.text
                return lastElement

        else:
            print "\t Warning: Heading number of heading (" + styName + ") is NaN"


# def doTable(par, t):
# 	lst = etree.SubElement(lastElement,"list")
# 	lst.set("rend","table")
# 	for r in t.rows:
# 		row = etree.SubElement(lst, "item")
# 		for c in r.cells:			
# 			cell = etree.SubElement(row, "rs")
# 			iterateRange(par, cell)
# 	return lastElement

def doNewList(styName, lastElement, cur_list_level):
    lst = etree.SubElement(lastElement, "list")

    if "List Bullet" in styName:
        lst.set("rend", "bullet")
        lst.set("n", str(cur_list_level))
    elif "List Number" in styName:
        lst.set("rend", "1")
        lst.set("n", str(cur_list_level))
    else:
        lst.set("n", str(cur_list_level))

    return lst


def closeStyle(styName, lastElement):
    global listOpen, lgOpen, citOpen, nestedCitOpen, speechOpen, nestedSpeechOpen, global_list_level

    flag = False

    # Lists
    if listOpen and "List" not in styName:
        listOpen = False
        flag = True
        while global_list_level > 1:
            global_list_level -= 1
            lastElement = lastElement.getparent().getparent()
        global_list_level = 0
        # this is for Citation List Bullet/Number (must pop out twice b/c of <quote> & <list>)
        # check if this breaks or works for citations where there is a citation paragraph/verse in a regular list
        if citOpen and "Citation" not in styName:
            citOpen = False
            return lastElement.getparent().getparent()
        else:
            lastElement = lastElement.getparent()

    # Close Citation / Nested Citation
    if citOpen and "Citation" not in styName:
        citOpen = False
        if nestedCitOpen and "Nested" not in styName:
            nestedCitOpen = False
            if lgOpen:
                lgOpen = False
                return lastElement.getparent().getparent().getparent()
            else:
                return lastElement.getparent().getparent()
        else:
            if lgOpen:
                lgOpen = False
                return lastElement.getparent().getparent()
            else:
                return lastElement.getparent()
    elif nestedCitOpen and "Nested" not in styName:
        nestedCitOpen = False
        if lgOpen:
            return lastElement.getparent().getparent()
        else:
            return lastElement.getparent()

    # Close Speech / Nested Speech
    if speechOpen and "Speech" not in styName:
        speechOpen = False
        if nestedSpeechOpen and "Nested" not in styName:
            nestedSpeechOpen = False
            if lgOpen:
                lgOpen = False
                return lastElement.getparent().getparent().getparent()
            else:
                return lastElement.getparent().getparent()
        else:
            return lastElement.getparent()
            # if lgOpen:
            #     lgOpen = False
            #     return lastElement.getparent().getparent()
            # else:
            #     return lastElement.getparent()
    elif nestedSpeechOpen and "Nested" not in styName:
        nestedSpeechOpen = False
        if lgOpen:
            return lastElement.getparent().getparent()
        else:
            return lastElement.getparent()

    # Close Verse 1 / Verse 2
    if lgOpen and not (
            "Verse 2" in styName or "Nested 1" in styName or "Nested 2" in styName or "1 Nested" in styName or "2 Nested" in styName):
        lgOpen = False
        return lastElement.getparent()
    else:
        return lastElement


def doParaStyles(par, prevSty, lastElement):
    # The main function called for paragraphs that are not headers
    # Sets the xml element to use and then calls iterate range
    global listOpen, lgOpen, citOpen, nestedCitOpen, speechOpen, nestedSpeechOpen, global_list_level

    styName = par.style.name
    # print styName
    if styName in ("Paragraph", "Paragraph Continued", "ParagraphContinued",  "Normal") and u'北宋' in par.text:
        print "In the paragraph"

    lastElement = closeStyle(styName, lastElement)

    #  Do "regular" paragraphs (just <p> tags)
    if styName in ("Paragraph", "Paragraph Continued", "ParagraphContinued",  "Normal"):
        p = etree.SubElement(lastElement, "p")
        iterateRange(par, p)
        return lastElement

    # fix this based on what Than/David say
    elif "Paragraph Single Spaced" == styName:
        # lastElement = lastElement.getparent()
        p = etree.SubElement(lastElement, "p")
        p.set('rend', 'nospace')
        iterateRange(par, p)
        return lastElement

    # textual citations
    elif "Citation" in styName:
        if not citOpen:
            lastElement = etree.SubElement(lastElement, "quote")
            citOpen = True

        # bulletted list in a citation
        if "Citation List Bullet" == styName or "List Bullet Citation" == styName:
            if "Citation List Bullet" not in prevSty:
                lastElement = etree.SubElement(lastElement, "list")
                lastElement.set("rend", "bullet")
                listOpen = True
            item = etree.SubElement(lastElement, "item")
            iterateRange(par, item)
            return lastElement

        # numbered list in a citation
        elif "Citation List Number" == styName or "List Number Citation" == styName:
            if "Citation List Number" not in prevSty:
                lastElement = etree.SubElement(lastElement, "list")
                lastElement.set("rend", "1")
                lastElement.set("n", "1")
                listOpen = True
            item = etree.SubElement(lastElement, "item")
            iterateRange(par, item)
            return lastElement


        # citOpen at starts of clause takes care of inserting <quote> when any new quote is started...so Paragraph Citation & Paragraph Citation Continued have same behavior
        elif "Z-Depracated Paragraph Citation" == styName or "Paragraph Citation" == styName:
            p = etree.SubElement(lastElement, "p")
            iterateRange(par, p)
            return lastElement
        elif "Citation Prose 2" == styName or "Paragraph Citation Continued" == styName:
            p = etree.SubElement(lastElement, "p")
            iterateRange(par, p)
            return lastElement

        elif "Paragraph Citation Nested" == styName:
            if "Paragraph Citation" not in prevSty and "Citation Prose 2" not in prevSty and "Z-Depracated Paragraph Citation" not in prevSty:
                print "\t Warning (IMPROPER CITATION NESTING): " + styName + " must be preceded by Paragraph Citation"
            nestedCitOpen = True
            quote = etree.SubElement(lastElement, "quote")
            iterateRange(par, quote)
            return quote

        elif "Citation Verse 1" == styName or "Verse Citation 1" == styName:
            lg = etree.SubElement(lastElement, "lg")
            l = etree.SubElement(lg, "l")
            iterateRange(par, l)
            lgOpen = True
            return lg

        elif "Citation Verse 2" == styName or "Verse Citation 2" == styName:
            l = etree.SubElement(lastElement, "l")
            iterateRange(par, l)
            return lastElement

        elif "Citation Verse Nested 1" == styName or "Verse Citation Nested 1" == styName:
            if "Citation Verse" not in prevSty and "Verse Citation" not in prevSty:
                print "\t Warning (IMPROPER CITATION NESTING): " + styName + " must be preceded by a Verse Citation"
            nestedCitOpen = True
            lg = etree.SubElement(lastElement, "lg")
            l = etree.SubElement(lg, "l")
            iterateRange(par, l)
            lgOpen = True
            return lg

        elif "Citation Verse Nested 2" == styName or "Verse Citation Nested 2" == styName:
            l = etree.SubElement(lastElement, "l")
            iterateRange(par, l)
            return lastElement

        else:
            print "\t Warning (Paragraph Style): " + styName + " is not a supported Citation Style"

    elif "List" in styName:
        listOpen = True
        try:
            cur_list_level = int(styName.split()[-1])
        except ValueError:
            cur_list_level = 1

        if cur_list_level - 1 > global_list_level:
            print "\t Warning (IMPROPER LIST NESTING): Jumped from List " + str(global_list_level) + " to List " + str(
                cur_list_level)
            print "\t\t List text: " + par.text

        # push new nested list level (or first level)
        if cur_list_level > global_list_level:
            while cur_list_level != global_list_level:
                if cur_list_level > 1:
                    lastElement = etree.SubElement(lastElement, "item")
                global_list_level += 1
                lastElement = doNewList(styName, lastElement, cur_list_level)
        # pop out a level or remain on same level
        else:
            while cur_list_level != global_list_level:
                global_list_level -= 1
                lastElement = lastElement.getparent().getparent()
        item = etree.SubElement(lastElement, "item")
        iterateRange(par, item)
        return lastElement


    elif "Speech" in styName:

        if "Speech Inline" == styName:
            q = etree.SubElement(lastElement, "q")
            iterateRange(par, q)
            return lastElement

        if not speechOpen:
            lastElement = etree.SubElement(lastElement, "q")
            speechOpen = True

        if "Speech Prose" == styName or "Speech Paragraph" == styName:
            p = etree.SubElement(lastElement, "p")
            iterateRange(par, p)
            return lastElement

        elif "Speech Prose Nested" == styName or "Speech Paragraph Nested" == styName:
            if "Speech Prose" not in prevSty and "Speech Paragraph" not in prevSty:
                print "\t Warning (IMPROPER SPEECH NESTING): " + styName + "must be preceded by Speech Paragraph"
            nestedSpeechOpen = True
            q = etree.SubElement(lastElement, "q")
            iterateRange(par, q)
            return q

        elif "Speech Verse 1" == styName:
            lg = etree.SubElement(lastElement, "lg")
            l = etree.SubElement(lg, "l")
            iterateRange(par, l)
            lgOpen = True
            return lg

        elif "Speech Verse 2" == styName:
            l = etree.SubElement(lastElement, "l")
            iterateRange(par, l)
            return lastElement

        elif "Speech Verse 1 Nested" == styName:
            if "Speech Verse" not in prevSty:
                print "\t Warning (IMPROPER SPEECH NESTING): " + styName + " must be preceded by a Speech Verse"
            nestedSpeechOpen = True
            lg = etree.SubElement(lastElement, "lg")
            l = etree.SubElement(lg, "l")
            iterateRange(par, l)
            lgOpen = True
            return lg

        elif "Speech Verse 2 Nested" == styName:
            l = etree.SubElement(lastElement, "l")
            iterateRange(par, l)
            return lastElement

        else:
            print "\t Warning (Paragraph Style): " + styName + " is not a supported Speech Style"
            return lastElement

    elif "Verse" in styName:
        if "Verse 1" == styName:
            lgOpen = True
            lg = etree.SubElement(lastElement, "lg")
            l = etree.SubElement(lg, "l")
            iterateRange(par, l)
            return lg

        elif "Verse 2" == styName:
            l = etree.SubElement(lastElement, "l")
            iterateRange(par, l)
            return lastElement

        else:
            print "\t Warning (Paragraph Style): " + styName + " is not a supported Verse Style"
            return lastElement



    elif "Section" in styName:
        ms = etree.SubElement(lastElement, "milestone")
        if "Chapter Element" in styName:
            ms.set("unit", "cle")
        elif "Top Level" in styName or "Division" in styName or "Level" in styName:
            ms.set("unit", "section")
            if "Second" in styName:
                ms.set("n", "2")
            elif "Third" in styName:
                ms.set("n", "3")
            elif "Fourth" in styName:
                ms.set("n", "4")
            else:
                ms.set("n", "1")
        else:
            print "\t Warning (Paragraph Style): " + styName + " is not a supported Section style"
        ms.set("rend", par.text)
        return lastElement



    else:
        print "\t Warning (Paragraph Style): " + styName + " is not supported"
        p = etree.SubElement(lastElement, "p")
        iterateRange(par, p)
        return lastElement


def interateRuns(par, lastElement):
    global footenotes, endnotes, footnoteNum, endnoteNum

    # styName is the current paragraph style
    styName = par.style.name
    runs = par.runs

    # empty paragraph (blank line)
    if len(runs) == 0:
        lastElement.text = " "
        return

    # iterate through runs in current paragraph.
    # Because we have run mergeRuns on document prior, each represents a new char style
    for run in runs:
        # charStyle is the current character style.
        charStyle = run.style.name

        # if character style of current run is same as current paragraph style
        if charStyle == "Default Paragraph Font":
            try:
                try:
                    elem.tail += run.text
                except TypeError:
                    elem.tail = run.text
            except (UnboundLocalError, AttributeError):
                try:
                    lastElement.text += run.text
                except TypeError:
                    lastElement.text = run.text

        elif charStyle == "Hyperlink":
            xref = etree.SubElement(lastElement, "xref")  # check if this prints hyperlink
            xref.set("n", run.text)

        else:
            # Page Number Print / Page Number
            if "Page Number" in charStyle or "PageNumber" == charStyle:
                # If there's a new opening bracket and the elem (milestone) already has a value. It's a new one being opened.
                msstrs = re.findall(r'\[(?:page ?:)?[^\]]+\]')
                for msstr in msstrs:
                    msdata = parseMilestoneText(msstr)
                if "[" in run.text:
                    try:
                        if len(elem.get("n")) > 0 and elem.get("n") != "":
                            oldelem = elem
                            elem = etree.SubElement(lastElement, "milestone")
                            elem.set("unit", "page")
                            elem.set("n", "")
                    except:
                        pass

                temp = run.text.replace(']/[', '').replace("[", "").replace("]", "").strip()
                if prevCharStyle == charStyle:
                    prev = elem.get("n")
                    elem.set("n", prev + temp)
                    currn = elem.get("n")
                    if elem.get('ed') is None and '-' in currn:
                        pts = currn.split('-')
                        elem.set("ed", pts[0])
                        elem.set("n", pts[1])
                else:
                    # Remove blank milestones if they occur
                    try:
                        n = elem.get('n')
                        if n == '':
                            elem.getparent().remove(elem)
                    except:
                        pass
                    elem = etree.SubElement(lastElement, "milestone")
                    elem.set("unit", "page")
                    if "-" in temp:
                        temp = temp.split("-")
                        elem.set("n", temp[1])
                        elem.set("ed", temp[0])
                    else:
                        elem.set("n", temp)
                    if "Page Number" == charStyle:
                        elem.set("rend", "digital")

            # Line Number Print / Line Number
            elif "Line Number" in charStyle or "TibLineNumber" == charStyle:
                temp = run.text.replace("line", "").replace("[", "").replace("]", "").strip()
                if prevCharStyle == charStyle:
                    prev = elem.get("n")
                    elem.set("n", prev + temp)
                else:
                    elem = etree.SubElement(lastElement, "milestone")
                    elem.set("unit", "line")
                    elem.set("n", temp)
                    if "Line Number" == charStyle:
                        elem.set("rend", "digital")

            elif charStyle == "Illegible":
                elem = etree.SubElement(lastElement, "gap")
                elem.set("n", run.text.split("[")[1].split("]")[0])
                elem.set("reason", "illegible")

            # continue with same character style within same XML tag
            elif charStyle == prevCharStyle:
                try:
                    try:
                        elem.text += run.text
                    except TypeError:
                        elem.text = run.text
                except (UnboundLocalError, AttributeError):
                    try:
                        lastElement.text += run.text
                    except TypeError:
                        lastElement.text = run.text

            # new character style or no character style
            else:
                elem = getElement(charStyle, lastElement)
                if elem == "none type":
                    try:
                        lastElement.text += run.text
                    except TypeError:
                        lastElement.text = run.text
                else:
                    if "footnote" in charStyle.lower():
                        elem.text = footnotes[footnoteNum - 1]
                    elif "endnote" in charStyle.lower():
                        elem.text = endnotes[endnoteNum - 1]
                    else:
                        elem.text = run.text
            prevCharStyle = charStyle

def iterateRange(par, lastElement):
    '''
    Iterates through the range of a paragrah.
    '''
    global footenotes, endnotes, footnoteNum, endnoteNum

    grandParent = lastElement   # record initial last element to return to
    # styName is the current paragraph style
    styName = par.style.name
    runs = par.runs

    # empty paragraph (blank line)
    if len(runs) == 0:
        lastElement.text = " "
        return

    # prevCharStyle is the most recent character style used. It is initialized to the current paragraph style.
    # prevCharStyle = styName  # This is not necessary since runs of the same charStyle have been concatenated
    elem = None

    # iterate through runs in current paragraph (all consquetive runs of the same style must be concatenated)
    for run in runs:
        # charStyle is the current character style.
        charStyle = run.style.name

        # if the character style is Default Paragraph
        if charStyle == "Default Paragraph Font":
            # If not currently in it, return to the grandparent (initial last element) as last Element
            if lastElement != grandParent:
                 elem = lastElement
                 lastElement = grandParent

            try:
                try:
                    elem.tail += run.text
                except TypeError:
                    elem.tail = run.text
            except (UnboundLocalError, AttributeError):
                try:
                    lastElement.text += run.text
                except TypeError:
                    lastElement.text = run.text

        elif charStyle == "Hyperlink":
            xref = etree.SubElement(lastElement, "xref")  # check if this prints hyperlink
            xref.set("n", run.text)

        # Page Number Print / Page Number or Line Numbers
        elif "Page Number" in charStyle or "PageNumber" == charStyle or \
                    "Line Number" in charStyle or "TibLineNumber" == charStyle:
                if charStyle == 'Line Number Digital' or charStyle == 'Page Number Digital':
                    elem = etree.SubElement(lastElement, "milestone")
                    unit = "page" if "Page" in charStyle else "line"
                    elem.set("unit", unit)
                    elem.set("rend", "digital")
                    elem.set("n", run.text)
                else:
                    msstrs = re.findall(r'\[[^\]]+\]', run.text)
                    for msstr in msstrs:
                        msdata = parseMilestoneText(msstr)
                        if len(msdata['n']) > 0:
                            elem = etree.SubElement(lastElement, "milestone")
                            elem.set("unit", msdata['unit'])
                            elem.set("n", msdata['n'])
                            if msdata['ed']:
                                elem.set("ed", msdata['ed'])

        elif charStyle == "Illegible":
            elem = etree.SubElement(lastElement, "gap")
            elem.set("n", run.text.split("[")[1].split("]")[0])
            elem.set("reason", "illegible")

        # Do Footnotes
        elif "footnote" in charStyle or "Footnote" in charStyle:
            #if type(lastElement) is etree._Element:
                #print u"last element: {0}, content : {1}, tail: {2}".format(lastElement.tag, lastElement.text, lastElement.tail)
            # if type(elem) is etree._Element:
            #     if type(elem.tail) is unicode and elem.tail[-1] == u'}':
            #         print u"Elem Var with critical: {1} | {0}".format(footnotes[footnoteNum - 1], elem.tail)
            critel = False

            if elem is not None and isinstance(elem.tail, (str, unicode)) and elem.tail[-1] == u'༽':
                critel = doCriticalElement(elem, 'tail')

            elif lastElement is not None and isinstance(lastElement.text, (str, unicode)) and lastElement.text[-1] == u'༽':
                critel = doCriticalElement(lastElement, 'text')

            else:
                elem = getElement(charStyle, lastElement)
                elem.text = footnotes[footnoteNum - 1]

            if critel is not None:
                elem = critel

        # Do Endnotes
        elif "endnote" in charStyle or "Endnote" in charStyle:
            elem = getElement(charStyle, lastElement)
            elem.text = endnotes[endnoteNum - 1]

        # Deal with tags that can contain other elements
        # If the nex charStyle matches the last recorded Element and there has been a child element
        # Then append this run to the tail of the child element
        elif matchesLastElement(charStyle, lastElement) and elem is not None:
            try:
                elem.tail += run.text
            except TypeError:
                elem.tail = run.text

        # If the new charstyle matches the parent of the last element, go up to it and add run as tail of last element
        elif matchesLastElement(charStyle, lastElement.getparent()):
            elem = lastElement
            lastElement = lastElement.getparent()
            try:
                elem.tail += run.text
            except TypeError:
                elem.tail = run.text

        # Otherwise call the getElement function to determine the element name based on charstyle and last element
        else:
            # if charStyle == "Name Personal Human" or charStyle == "Sa bcad":
            #     print charStyle, etree.tostring(lastElement)

            newel = getElement(charStyle, lastElement)
            if newel != 'none type':
                lastElement = newel             # Set this as last element so it contains any next elements until Default Paragraph Style is found
                lastElement.text = run.text
                processSpecialElements(lastElement, charStyle)
                elem = None
            else:                               # if char style returns nonetype then append to current elem tail or if not, then lastElem text
                if elem is not None:
                    try:
                        elem.tail += run.text
                    except TypeError:
                        elem.tail = run.text
                else:
                    try:
                        lastElement.text += run.text
                    except TypeError:
                        lastElement.text = run.text

            # ORiginal code
            # elem = getElement(charStyle, lastElement)
            # if elem == "none type":
            #     try:
            #         lastElement.text += run.text
            #     except TypeError:
            #         lastElement.text = run.text
            #
            # else:
            #     elem.text = run.text


def matchesLastElement(charStyle, lastel):
    '''
    Check to see if the current new tag is exactly the same as the last parent to determine how to deal with text

    :param charStyle:
    :return:
    '''
    # Get test element based on charstyle
    fauxparent = etree.fromstring('<faux></faux>') # a fake parent to call get element with
    testel = getElement(charStyle, fauxparent, False)

    # Check if tag is the same as last el
    if testel != "none type" and testel.tag == lastel.tag:
        # Go through test el's attributes and match to last ele
        for att in testel.attrib:
            if testel.get(att) != lastel.get(att):
                return False

        # Go through last el's attributes and match with test el
        for att in lastel.attrib:
            if lastel.get(att) != testel.get(att):
                return False
    else:
        return False

    return True  # if it makes it here they are identical


def processSpecialElements(el, chSty):
    '''
    Especially process the text of certain elements based on style name such as "Abbreviation"

    :param el:
    :param chSty:
    :return:
    '''
    #print "in process special: {}".format(chSty)
    if chSty == u'Abbreviation':
        eltxt = unicode(el.text)
        pts = eltxt.split(u'༼')
        if len(pts) > 1:
            el.text = pts[0]
            el.set('expan', pts[1].replace(u'༽',u''))


def doCriticalElement(elem, txttype='tail'):
    '''
    Creates the markup for a different reading in a critical edition of a text based on {} and footnotes.
    An example of the reading in the footnote: Lh: འདི་སྐད་ (1123a.4), KND: སྡེ་དགེ་ (56.3).
    :param elem: the element whose tail or text has the critical edition markup in it
    :return:
    '''

    global footnotes, footnoteNum, debugme



    if txttype == 'tail':
        txt = elem.tail
    elif txttype == 'text':
        txt = elem.text

    if debugme:
        print u"\n---------------------------"
        print u"type is: {}".format(txttype)
        print u"Elem is: {}".format(elem.tag)
        print u"text is: {}".format(txt)
        print u"fnum is: [{}]".format(footnoteNum)
        print u"footnote text is: {}".format(footnotes[footnoteNum])

    if not isinstance(txt, unicode):
        print u"Could not find text ({}) to build apparatus: {}".format(txttype, unicode(etree.tostring(elem)))

    txtpts = txt.replace(u'༼༽',u'༼none༽').split(u'༼')

    if len(txtpts) == 2:
        # parse the elements text (or tail) to find the part surrounded by ༼...༽ which is the lemma
        pretext = txtpts[0]
        #print u"Pretext: {}".format(pretext)
        lemma = txtpts[1].replace(u'༽','')

        # Get the corresponding footnote text (before increasing the number by 1). These are the readings
        rdgs = footnotes[footnoteNum]
        footnoteNum += 1  # increase the footnote for the next one

        # Split the individual readings on the semicolons. See: https://docs.google.com/document/d/11C_YYg6Y4JHVa2tdAd9Jqych-1PGK-dAUS2cNZ2xS1o/edit#heading=h.6ae0yzdsuvfs
        temp = rdgs.split(u';')
        rdgs = []
        for r in temp:
            rdg = parseReading(r.strip()) # parse each reading into a dictionary of wit(sigla), page, and text
            if debugme:
                print u"reading: {}".format(r)
                print u"reading dict: {}".format(rdg)
            if rdg:
                rdgs.append(rdg)

        # build the <app><lem></lem><rdg></rdg></app> element
        app = etree.Element("app")
        lem = etree.SubElement(app, "lem")
        if lemma == u'none':
            lem.set('rend', 'omits')
        else:
            lem.text = lemma
        for r in rdgs:
            rdg = etree.SubElement(app, "rdg")
            if 'text' in r:
                rdg.text = r['text']
            if 'page' in r:
                rdg.set('n', r['page'])
            if 'wit' in r:
                rdg.set('wit', r['wit'])

        if txttype == 'tail':
            if debugme:
                print u"pretxt in tail: {}".format(pretext)
                print u"elem in tail: {}".format(elem.tag)
            elem.tail = pretext
            epar = elem.getparent()
            epar.insert(epar.index(elem) + 1, app)
            return app

        elif txttype == 'text':
            elem.text = pretext
            elem.append(app)
            return app
        else:
            print u"Unknown text position type: {}".format(txttype)

    else:
        print u"Warning: Incorrect number of parts to split in critical element text: {}".format(txt)

    return False

def parseReading(rdgtxt):
    global footnoteNum

    #print u"readging text: {}".format(rdgtxt)
    rdg = {}
    pts = rdgtxt.split(u':')
    if len(pts) == 2:
        wits = pts[0].strip()
        wits = re.sub(u',', u' ', wits)
        wits = re.sub(u'\s+', u' ', wits)
        rdg['wit'] = wits
        pts = pts[1].replace(u')', '').split(u'(')
        rdg['text'] = pts[0].strip()
        if len(pts) > 1:
            rdg['page'] = pts[1].strip()
        return rdg
    else:
        print u"Warning: Reading in footnote #{} does not have colon: {}".format(footnoteNum, rdgtxt)
    return None


# implement fully
#  This is not called (ndg, 2018-08-20)
# def iterateNote(run, lastElement, styName):
#     # lastElement.text = run.text
#     # FIX/TEST THIS
#     prevCharStyle = styName
#
#     # iterate through characters in the footnote
#     for char in run.text:
#         charStyle = char.style.name
#         # char style is same as paragraph style
#         if charStyle == styName or charStyle == "Default Paragraph Font":
#             try:
#                 try:
#                     elem.tail += char
#                 except TypeError:
#                     elem.tail = char
#             except (UnboundLocalError, AttributeError):
#                 try:
#                     lastElement.text += char
#                 except TypeError:
#                     lastElement.text = char
#             prevCharStyle = styName
#
#         elif charStyle == "Hyperlink":
#             xref = etree.SubElement(lastElement, "xref")  # check if this prints hyperlink
#             xref.set("n", char)
#             # handle displaying link text next to tag
#             prevCharStyle = charStyle
#         else:
#             # continue with same character style
#             if charStyle == prevCharStyle:
#                 try:
#                     elem.text += char
#                 except TypeError:
#                     elem.text = char
#
#             # Page Number,digital
#             if "Page Number" == charStyle:
#                 elem = etree.SubElement(lastElement, "milestone")
#                 elem.set("unit", "page")
#                 temp = char
#                 if char[0] == "[":
#                     temp = temp[1:]
#                 if char[-1] == "]":
#                     temp = temp[:-1]
#                 elem.set("n", temp)
#                 elem.set("rend", "digital")
#
#             # Line Number,digital
#             elif "Line Number" == charStyle:
#                 elem = etree.SubElement(lastElement, "milestone")
#                 elem.set("unit", "line")
#                 temp = char
#                 if char[0] == "[":
#                     temp = temp[1:]
#                 if char[-1] == "]":
#                     temp = temp[:-1]
#                 elem.set("n", temp)
#                 elem.set("rend", "digital")
#
#             # Page Number Number Print Edition,pnp"
#             elif "Page Number Print" in charStyle or "PageNumber" == charStyle:
#                 temp = run.text.replace("page", "").replace("[", "").replace("]", "").strip()
#                 if len(temp) > 0:
#                     elem = etree.SubElement(lastElement, "milestone")
#                     elem.set("unit", "page")
#                     if "-" in temp:
#                         temp = temp.split("-")
#                         elem.set("n", temp[1])
#                         elem.set("ed", temp[0])
#                     else:
#                         elem.set("n", temp)
#
#             # Line Number Print,lnp
#             elif "Line Number Print" in charStyle or "TibLineNumber" == charStyle:
#                 temp = run.text.replace("line", "").replace("[", "").replace("]", "").strip()
#                 if len(temp) > 0:
#                     elem = etree.SubElement(lastElement, "milestone")
#                     elem.set("unit", "line")
#                     if "-" in temp:
#                         temp = temp.split("-")
#                         elem.set("n", temp[1])
#                         elem.set("ed", temp[0])
#                     else:
#                         elem.set("n", temp)
#
#             elif charStyle == "Illegible":
#                 elem = etree.SubElement(lastElement, "gap")
#                 elem.set("n", run.text.split("[")[1].split("]")[0])
#                 elem.set("reason", "illegible")
#
#             # set new character style
#             else:
#                 elem = getElement(charStyle, lastElement)
#                 if elem == "none type":
#                     try:
#                         lastElement.text += char
#                     except TypeError:
#                         lastElement.text = char
#                 else:
#                     elem.text = char
#             prevCharStyle = charStyle


def parseMilestoneText(msstr):
    '''
    A function to parse the text of a milestone phrase, e.g. "[page Ab1-45b]" or "[line Tr-89.1]" or just "[901b]"
    And return a dictionary of data to create the milestone tag

    :param msstr: the milestone string to parse includes the brackets []
    :return: a dictionary with the following keys:
                "unit" => page or line
                "n" => value of the n attribute, e.g. "45b", "89.1", "901b"
                "ed" => edition sigla, e.g. "Ab1", "Tr", or False
    '''
    data = {}
    msstr = msstr.replace('[', '').replace(']', '')
    # If the milestone string has a space the unit type is first
    if " " in msstr:
        pts = msstr.split(' ')
        data['unit'] = pts[0]
        msstr = pts[1]
    elif "." in msstr:
        data['unit'] = 'line'
    else:
        data['unit'] = 'page'

    msstr = msstr.replace(' ', '') # make sure all whitespace is gone

    if '-' in msstr:
        pts = msstr.split('-')
        data['ed'] = pts[0]
        data['n'] = pts[1]
    else:
        data['ed'] = False
        data['n'] = msstr

    return data

def getElement(chStyle, lastElement, warn=True):
    '''
    Gets the element based on the Word-style uses the "getStyleTagDef(stnm)" function in styleElement.py

    :param chStyle:
    :param lastElement:
    :param warn:
    :return:
    '''

    global footnoteNum, endnoteNum

    eldef = styEl.getStyleTagDef(chStyle)

    if "Footnote" in chStyle or "footnote" in chStyle:
        footnoteNum += 1
        elem = etree.SubElement(lastElement, "note")
        elem.set("n", str(footnoteNum))

    elif "Endnote" in chStyle or "endnote" in chStyle:
        endnoteNum += 1
        elem = etree.SubElement(lastElement, "note")
        elem.set("n", str(endnoteNum))

    elif eldef:
        elem = etree.SubElement(lastElement, eldef['tag'])
        if 'attributes' in eldef:
            atts = eldef['attributes']
            if isinstance(atts, dict):
                for nm, val in atts.iteritems():
                    elem.set(nm, val)

    else:
        global debugme, unsupported_char
        if debugme is True:
            print "\t Warning (Character Style): " + chStyle + " is not supported"

        if chStyle in unsupported_char:
            unsupported_char[chStyle] += 1
        else:
            unsupported_char[chStyle] = 1

        return "none type"

    return elem


def getNewLastElement(elname="p"):
    nle = etree.Element(elname)
    root.find('text').append(nle)
    return nle


def mergeRuns(doc):
    '''
    Take a document and go through all runs in all paragraphs, if two consecutive runs have the same style, then merge them

    :param doc:
    :return:
    '''
    for para in doc.paragraphs:
        runs2remove = []
        lastrun = False
        # Merge runs with same style
        for n, r in enumerate(para.runs):
            if lastrun is not False and r.style.name == lastrun.style.name:
                lastrun.text += r.text
                runs2remove.append(r)
            else:
                lastrun = r
        # Remove all runs thus merged
        for rr in runs2remove:
            el = rr._element
            el.getparent().remove(el)


def convertDoc(inputFile, outpath):
    global tableOpen, listOpen, lgOpen, citOpen, nestedCitOpen, speechOpen, nestedSpeechOpen, bodyOpen, backOpen, \
        frontOpen, global_header_level, inDocument, global_list_level, footnoteNum, endnoteNum, idTracker, titlePage, \
        root, badheader_text, convert_options

    if debugme:
        print "converting {0} to {1}\n".format(inputFile, outpath)

    # reset global variables
    tableOpen = listOpen = lgOpen = citOpen = nestedCitOpen = speechOpen = nestedSpeechOpen = False
    bodyOpen = backOpen = frontOpen = False
    global_header_level = 0
    inDocument = False
    global_list_level = 0
    footnoteNum = endnoteNum = 0
    idTracker = []
    titlePage = None

    # process foot/endnotes into lists
    print "Doing footnotes ..."
    doFootEndnotes(inputFile)

    # read input file
    document = Document(inputFile)

    # process metadata table
    print "Processing metadata table ..."
    try:
        metaTable = document.tables[0]

    except:
        print "\t Error: metatable not included"
        sys.exit(1)

    metaText = doMetadata(metaTable)

    # create lxml element tree from metadata info
    metaText = metaText.encode('utf-8')
    parser = etree.XMLParser(ns_clean=True, recover=True, encoding='utf-8')
    root = etree.fromstring(metaText, parser=parser)
    # print etree.tostring(root, encoding='UTF-8', xml_declaration=False)

    # Debugging
    # with open("../out/rootdoctest.xml", "wb") as fout:
    #     docType = "<!DOCTYPE TEI.2 SYSTEM \"http://www.thlib.org:8080/cocoon/texts/catalogs/xtib3.dtd\">"
    #     toString = etree.tostring(root, pretty_print=True, encoding='utf-8', xml_declaration=True, doctype=docType)
    #     fout.write(toString)
    # exit(0)

    # iterate through paragraphs
    lastElement = root.find('text')
    prevSty = ''
    mergeRuns(document) # Merge consequetive runs of the same style, so each run represents a new style
    print "Processing text ..."
    for par in document.paragraphs:
        # Do the Headers
        if "Heading" in par.style.name or "Interstitial Section" in par.style.name:
            # inDocument avoids including any paragraphs before the first structural heading
            inDocument = True
            lastElement = doHeaders(par, lastElement, root)
        # Set in Heading. See above
        elif inDocument:
            lastElement = doParaStyles(par, prevSty, lastElement)
        # Do the Title of the Word Doc
        elif "Title" == par.style.name:
            doTitle(par, lastElement)
        # Top level paragraph that is not Front, Body, or Back (Heading) is wrong. So warn
        else:
            if par.text:
                badheader_text.append("{}".format(par.text.strip(" \s\n\r")))

        prevSty = par.style.name

        if lastElement is None:
            print "\t No last element after processing paragraph: "
            print "\t\t Text: " + par.text
            lastElement = getNewLastElement()

    closingComments(lastElement)

    # Determine Name for Resulting XML file
    fname = inputFile.split("/")[-1].split(".")[0] + '.xml'
    fpth = os.path.join(outpath, fname)
    while os.path.isfile(fpth):
        userin = raw_input("The file {} already exists. Overwrite it (y/n/q): ".format(fname))
        if userin == 'y':
            break
        elif userin == 'n':
            fname = raw_input("Enter a new file name: ")
            fpth = os.path.join(outpath, fname)
        else:
            exit(0)

    # Write XML File
    with open(fpth, "wb") as outfile:
        dtdpath = convert_options.dtdpath
        docType = "<!DOCTYPE TEI.2 SYSTEM \"{}xtib3.dtd\">".format(dtdpath)
        toString = etree.tostring(root, pretty_print=True, encoding='utf-8', xml_declaration=True, doctype=docType)
        outfile.write(toString);


def getMetaFieldsFromTemplate():
    global metaTemplate
    # load metadata schema
    try:
        # if
        f = open(metaTemplate, 'rb')
        metaText = f.read()
    except:
        print "\t Error: teiHeader.dat file not in current working directory"
        print "\t Current directory: {0}".format(os.getcwd())
        sys.exit(1)

    allrepstrs = re.findall(r'\{[^\}]+\}', metaText, re.MULTILINE)
    replist = list()
    for repstr in allrepstrs:
        replist.append("{}".format(repstr.replace('{','').replace('}','')))

    replist.sort()
    return replist


########## MAIN ##########
def main():
    global metaTemplate, convert_options, unsupported_char, badheader_text, debugme

    # Generate the arg parser and options
    parser = argparse.ArgumentParser(description='Convert THL Word marked up documents to THL TEI XML',
                                     formatter_class=argparse.ArgumentDefaultsHelpFormatter)
    parser.add_argument('source', nargs='+', help='The space-separated paths to one or more Word documents to be converted. (Paths can be relative.)')
    parser.add_argument('-o', '--out', default='../out', help='The relative path to the outfolder')
    parser.add_argument('-mtf', '--metafields', action='store_true', help='List the metadata fields in the template')
    parser.add_argument('-t', '--template', default='teiHeader.dat', help='Relative path to a metadata table XML template')
    parser.add_argument('-dtd', '--dtdpath', default='http://www.thlib.org:8080/cocoon/texts/catalogs/', help='Path to the xtib3.dtd to add to the xmlfile')
    args = parser.parse_args()

    # Deal with template argument (by default the global variable metaTemplate is set to "teiHeader.dat"
    if not args.template:  # Default Template
        print "Using default template teiHeider.dat"

    elif os.path.isfile(args.template): # Template from File
        if debugme:
            print "template arg is: {}".format(args.template)
        metaTemplate = args.template

    else:  # Throw Error is not Template given
        print "Error: The Template path you supplied is not valid"
        exit(0)

    # if -mtf or metafields argument is chosen, list metadata fields in the template
    if args.metafields:
        replist = getMetaFieldsFromTemplate()
        for item in replist:
            print item
        exit(0)

    # Check that outpath is valid
    if not os.path.isdir(args.out):
        print "Error: The destination directory for the XML files (outpath) is not a valid directory"
        exit(0)

    # if No source given
    if not args.source or (isinstance(args.source, list) and len(args.source) == 0):
        print "Error: You need to supply a directory or file name to convert!"
        parser.print_help()
        exit(0)

    # Convert source path to list of relative file paths
    if os.path.isdir(args.source[0]):
        source_path = args.source[0]
        new_list = list()
        for sfile in os.listdir(source_path):
            if sfile.endswith(".docx"):
                new_list.append(os.path.join(source_path, sfile))
        if len(new_list) > 0:
            args.source = new_list
        else:
            print "Error: No valid files found in path given. All files must be of extension *.docx"
            exit(0)

    if debugme:
        print "{}".format(args) # for debugging for time being

    convert_options = args # save options globally just in case needed somewhere
    mysuccess = False

    for cfile in args.source:
        print "**********************************"
        print "Converting {} to XML...".format(cfile)
        convertDoc(cfile, convert_options.out)
        mysuccess = True

    print "Done!"

    if mysuccess:
        if badheader_text:
            print "\nThe following paragraphs were improperly nested (outside front, body, or back):"
            for btxt in badheader_text:
                print "\t{}".format(btxt)

        if unsupported_char:
            print "\nThe following character styles were not supported: "
            for styl, numf in unsupported_char.iteritems():
                print "\t{} ({} times)".format(styl, numf)

        print "\nConversion successful!"

    else:
        print "\nConversion failed!"

    print "***********************************"
    # Old code
    # docs = []
    # inpath = False
    # outpath = os.path.join(os.getcwd(), '../out/')
    # getOutPath = False
    # for item in sys.argv[1:]:
    #     if getOutPath:
    #         outpath = os.path.join(os.getcwd(), item)
    #         getOutPath = False
    #     elif item == '-o':
    #         getOutPath = True
    #     elif item.endswith(".docx"):
    #         docs.append(item)
    #     else:
    #         fullpath = os.path.join(os.getcwd(), item)
    #         if os.path.isdir(fullpath):
    #             inpath = fullpath
    #         else:
    #             "\t Warning (IMPROPER ARGUMENT): " + item + " is not a docx file in the current working directory"
    #
    # if inpath:
    #     status = "is" if os.path.isdir(inpath) else "is not"
    #     print "Converting all .docx in the path: {0}  (It {1} a directory)".format(inpath, status)
    # elif len(docs) > 0:
    #     if debugme:
    #         print "Converting the following docs: {0}".format(', '.join(docs))
    # else:
    #     print "\tWarnging (INCORRECT ARGUMENTS): Neither docs not inpath given"
    #     exit(0)
    #
    # status = "is" if os.path.isdir(outpath) else "is not"
    # if debugme:
    #     print "Outpath for the converted xml is: {0} (It {1} a directory)".format(outpath, status)
    #
    # mysuccess = False
    # # Process all .docx files in inpath
    # if inpath:
    #     for item in os.listdir(inpath):
    #         currentPath = os.path.join(inpath, item)
    #         if item.endswith(".docx") and os.path.isfile(currentPath):
    #             print "Converting " + item + " to XML..."
    #             convertDoc(currentPath, outpath)
    #             mysuccess = True
    #
    # # Process list of files given as parameters
    # else:
    #     for item in docs:
    #         currentPath = os.path.join(os.getcwd(), item)
    #         if item.endswith(".docx") and os.path.isfile(currentPath):
    #             print "Converting " + item + " to XML..."
    #             convertDoc(currentPath, outpath)
    #             mysuccess = True
    #
    # if mysuccess:
    #     if badheader_text:
    #         print "\n\tThe following paragraphs were improperly nested (outside front, body, or back):"
    #         for btxt in badheader_text:
    #             print "\t\t{}".format(btxt)
    #
    #     if unsupported_char:
    #         print "\n\tThe following character styles were not supported: "
    #         for styl, numf in unsupported_char.iteritems():
    #             print "\t\t{} ({} times)".format(styl, numf)
    #
    #     print "\nConversion successful!"

main()
